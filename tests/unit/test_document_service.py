import shutil
from pathlib import Path

import fitz
import pytest
from docx import Document

from app.services.document_service import DocumentService


FIXTURE_TMP_DIR = Path(__file__).resolve().parents[1] / "fixtures" / "document_service_tmp"


@pytest.fixture
def document_tmp_dir():
    FIXTURE_TMP_DIR.mkdir(parents=True, exist_ok=True)
    yield FIXTURE_TMP_DIR
    shutil.rmtree(FIXTURE_TMP_DIR, ignore_errors=True)


def test_extract_text_from_pdf(document_tmp_dir):
    pdf_path = document_tmp_dir / "resume.pdf"
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "Jane Candidate\nPython Developer")
    document.save(pdf_path)
    document.close()

    text = DocumentService().extract_text(str(pdf_path))

    assert "Jane Candidate" in text
    assert "Python Developer" in text


def test_extract_text_from_docx(document_tmp_dir):
    docx_path = document_tmp_dir / "resume.docx"
    document = Document()
    document.add_paragraph("Jane Candidate")
    document.add_paragraph("Python Developer")
    document.save(docx_path)

    text = DocumentService().extract_text(str(docx_path))

    assert text == "Jane Candidate\n\nPython Developer"


def test_extract_text_rejects_unsupported_extension(document_tmp_dir):
    text_path = document_tmp_dir / "resume.txt"
    text_path.write_text("plain text resume", encoding="utf-8")

    with pytest.raises(ValueError, match="Unsupported document type"):
        DocumentService().extract_text(str(text_path))


def test_extract_text_rejects_missing_file(document_tmp_dir):
    missing_path = document_tmp_dir / "missing.pdf"

    with pytest.raises(FileNotFoundError, match="Document not found"):
        DocumentService().extract_text(str(missing_path))
